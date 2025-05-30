# -*- coding: utf-8 -*-
"""
文档处理模块
负责读取、解析和写入Word文档，以及应用排版格式。
"""

import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from ..utils.logger import app_logger
from ..utils.file_utils import generate_output_filename, is_valid_docx, backup_file
from ..utils.font_manager import FontManager

class DocProcessor:
    """文档处理器，负责读取、解析和写入Word文档"""
    
    def __init__(self):
        """初始化文档处理器"""
        self.document = None
        self.input_file = None
        self.output_file = None
        self.paragraphs_text = []
        
        # 使用字体管理器获取字体信息
        self.font_manager = FontManager()
        
        # 字号映射
        self.font_size_mapping = {
            "小二": Pt(18),
            "三号": Pt(16),
            "小三": Pt(15),
            "四号": Pt(14),
            "小四": Pt(12),
            "五号": Pt(10.5),
            "小五": Pt(9),
            "六号": Pt(7.5)
        }
        
        app_logger.debug(f"文档处理器初始化完成，字号映射: {list(self.font_size_mapping.keys())}")
    
    def read_document(self, file_path):
        """
        读取Word文档内容
        
        Args:
            file_path: 文档路径
            
        Returns:
            是否成功读取文档
        """
        if not is_valid_docx(file_path):
            app_logger.error(f"无效的Word文档: {file_path}")
            return False
        
        try:
            self.document = Document(file_path)
            self.input_file = file_path
            self.paragraphs_text = [p.text for p in self.document.paragraphs]
            
            app_logger.info(f"成功读取文档: {file_path}")
            app_logger.debug(f"文档包含 {len(self.document.paragraphs)} 个段落")
            return True
        except Exception as e:
            app_logger.error(f"读取文档失败: {file_path}, 错误: {str(e)}")
            return False
    
    def get_document_text(self):
        """
        获取文档的纯文本内容
        
        Returns:
            文档的段落文本列表
        """
        if not self.document:
            app_logger.warning("尚未加载文档")
            return []
        
        return self.paragraphs_text
    
    def apply_formatting(self, formatting_instructions, custom_save_path=None):
        """
        根据排版指令应用格式
        
        Args:
            formatting_instructions: 排版指令，包含元素类型和格式信息
            custom_save_path: 自定义保存路径，如果指定则使用该路径
            
        Returns:
            是否成功应用格式
        """
        if not self.document:
            app_logger.error("尚未加载文档，无法应用格式")
            return False
        
        try:
            # 记录格式化指令，方便调试
            app_logger.debug(f"格式化指令: {formatting_instructions}")
            
            # 检查formatting_instructions结构
            if not isinstance(formatting_instructions, dict):
                app_logger.error(f"排版指令不是字典类型: {type(formatting_instructions)}")
                return False
            
            elements = formatting_instructions.get('elements', [])
            if not isinstance(elements, list):
                app_logger.error(f"元素列表不是列表类型: {type(elements)}")
                return False
            
            app_logger.info(f"开始应用排版格式，共有 {len(elements)} 个元素")
            
            # 记录自定义保存路径
            if custom_save_path:
                app_logger.info(f"使用自定义保存路径: {custom_save_path}")
            
            # 备份原始文档
            backup_path = backup_file(self.input_file)
            app_logger.info(f"文件备份成功: {self.input_file} -> {backup_path}")
            
            # 创建新文档以应用格式
            try:
                new_doc = Document()
                app_logger.debug("成功创建新文档")
            except Exception as e:
                app_logger.error(f"创建新文档失败: {str(e)}")
                return False
            
            # 处理每个元素
            for i, element in enumerate(elements):
                try:
                    app_logger.debug(f"处理第 {i+1}/{len(elements)} 个元素")
                    self._process_element(new_doc, element)
                except Exception as e:
                    app_logger.error(f"处理第 {i+1} 个元素时发生错误: {str(e)}")
                    # 继续处理下一个元素，不中断整个过程
            
            # 生成输出文件名
            try:
                self.output_file = generate_output_filename(self.input_file, custom_save_path)
                app_logger.debug(f"生成输出文件名: {self.output_file}")
            except Exception as e:
                app_logger.error(f"生成输出文件名失败: {str(e)}")
                # 如果指定了自定义保存路径，则使用该路径
                if custom_save_path and os.path.isdir(custom_save_path):
                    filename = os.path.basename(self.input_file).replace('.docx', '_已排版.docx')
                    self.output_file = os.path.join(custom_save_path, filename)
                else:
                    self.output_file = self.input_file.replace('.docx', '_已排版.docx')
                app_logger.debug(f"使用默认输出文件名: {self.output_file}")
            
            # 保存文档
            try:
                # 确保输出目录存在
                output_dir = os.path.dirname(self.output_file)
                if not os.path.exists(output_dir):
                    os.makedirs(output_dir)
                    app_logger.debug(f"创建输出目录: {output_dir}")
                
                new_doc.save(self.output_file)
                app_logger.info(f"成功应用排版格式并保存到: {self.output_file}")
                return True
            except Exception as e:
                app_logger.error(f"保存文档失败: {str(e)}")
                return False
        except Exception as e:
            app_logger.error(f"应用排版格式失败: {str(e)}")
            import traceback
            app_logger.error(f"详细错误信息: {traceback.format_exc()}")
            return False
    
    def _process_element(self, doc, element):
        """
        处理单个排版元素
        
        Args:
            doc: 目标文档对象
            element: 排版元素信息
        """
        try:
            # 记录当前处理的元素信息
            app_logger.debug(f"开始处理元素: {element}")
            
            # 检查元素结构
            if not isinstance(element, dict):
                app_logger.error(f"元素不是字典类型: {type(element)}")
                return
            
            content = element.get('content', '')
            app_logger.debug(f"元素内容: {content[:50]}...")
            
            element_type = element.get('type', '正文')
            app_logger.debug(f"元素类型: {element_type}")
            
            format_info = element.get('format', {})
            app_logger.debug(f"格式信息: {format_info}")
            
            # 检查format_info结构
            if not isinstance(format_info, dict):
                app_logger.error(f"格式信息不是字典类型: {type(format_info)}")
                format_info = {}
            
            # 添加段落
            paragraph = doc.add_paragraph()
            app_logger.debug("成功添加段落")
            
            run = paragraph.add_run(content)
            app_logger.debug("成功添加文本运行")
            
            # 应用字体
            self._apply_font(run, format_info)
            app_logger.debug("成功应用字体格式")
            
            # 应用段落格式
            self._apply_paragraph_format(paragraph, format_info, element_type)
            app_logger.debug("成功应用段落格式")
            
            app_logger.debug(f"完成元素处理: {element_type}, 内容: {content[:20]}...")
        except Exception as e:
            app_logger.error(f"处理元素时发生异常: {str(e)}")
            # 不抛出异常，继续处理下一个元素
    
    def _apply_font(self, run, format_info):
        """
        应用字体格式
        
        Args:
            run: 文本运行对象
            format_info: 格式信息
        """
        try:
            # 字体名称
            font_name = format_info.get('font', '宋体')
            app_logger.debug(f"原始字体名称: {font_name}")
            
            # 获取映射后的字体名称，直接使用新方法，不检查可用性
            document_font = self.font_manager.get_font_for_document(font_name)
            app_logger.debug(f"用于文档的字体名称: {document_font}")
            
            # 设置英文字体名称
            run.font.name = document_font
            
            # 设置中文字体（使用相同的字体）
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), document_font)
                app_logger.debug(f"成功设置中文字体: {document_font}")
            except Exception as e:
                app_logger.error(f"设置中文字体失败: {str(e)}")
                
            # 字体大小
            font_size = format_info.get('size', '五号')
            app_logger.debug(f"原始字体大小: {font_size}")
            
            if isinstance(font_size, str):
                mapped_size = self.font_size_mapping.get(font_size, Pt(10.5))  # 默认五号字体
                app_logger.debug(f"映射后的字体大小: {mapped_size}")
                font_size = mapped_size
            
            run.font.size = font_size
            
            # 粗体
            bold = format_info.get('bold', False)
            app_logger.debug(f"设置粗体: {bold}")
            run.bold = bold
            
            # 斜体
            italic = format_info.get('italic', False)
            app_logger.debug(f"设置斜体: {italic}")
            run.italic = italic
            
            # 下划线
            underline = format_info.get('underline', False)
            app_logger.debug(f"设置下划线: {underline}")
            run.underline = underline
            
        except Exception as e:
            app_logger.error(f"应用字体格式时发生异常: {str(e)}")
            # 不抛出异常，使用默认字体设置
    
    def _apply_paragraph_format(self, paragraph, format_info, element_type):
        """
        应用段落格式
        
        Args:
            paragraph: 段落对象
            format_info: 格式信息
            element_type: 元素类型
        """
        try:
            # 行间距
            line_spacing = format_info.get('line_spacing', 1.5)
            app_logger.debug(f"设置行间距: {line_spacing}")
            
            if isinstance(line_spacing, (int, float)):
                try:
                    # 防止异常大的行间距值导致程序崩溃
                    if line_spacing > 36.0 or line_spacing < 0.5:
                        app_logger.warning(f"检测到异常行间距值: {line_spacing}，将使用默认值1.5")
                        line_spacing = 1.5
                    
                    if line_spacing == 1.0:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        app_logger.debug("设置单倍行间距")
                    elif line_spacing == 1.5:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        app_logger.debug("设置1.5倍行间距")
                    elif line_spacing == 2.0:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                        app_logger.debug("设置双倍行间距")
                    else:
                        # 自定义行间距，安全值
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        paragraph.paragraph_format.line_spacing = Pt(line_spacing)
                        app_logger.debug(f"设置自定义行间距: {line_spacing}磅")
                except Exception as e:
                    app_logger.error(f"设置行间距失败: {str(e)}")
            
            # 对齐方式
            alignment = format_info.get('alignment', 'left')
            app_logger.debug(f"设置对齐方式: {alignment}")
            
            try:
                if alignment == 'center':
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif alignment == 'right':
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif alignment == 'justify':
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                else:  # 默认左对齐
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            except Exception as e:
                app_logger.error(f"设置对齐方式失败: {str(e)}")
            
            # 首行缩进
            try:
                first_line_indent = format_info.get('first_line_indent', None)
                app_logger.debug(f"首行缩进设置: {first_line_indent}")
                
                if first_line_indent is not None:
                    if isinstance(first_line_indent, (int, float)):
                        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent)
                        app_logger.debug(f"设置自定义首行缩进: {first_line_indent}磅")
                    else:
                        app_logger.warning(f"首行缩进值类型错误: {type(first_line_indent)}")
                elif element_type == '正文':  # 正文默认缩进2个字符
                    paragraph.paragraph_format.first_line_indent = Pt(21)  # 约2个中文字符
                    app_logger.debug("设置正文默认首行缩进: 21磅")
            except Exception as e:
                app_logger.error(f"设置首行缩进失败: {str(e)}")
            
            # 段前间距
            try:
                before_spacing = format_info.get('before_spacing', None)
                app_logger.debug(f"段前间距设置: {before_spacing}")
                
                if before_spacing is not None:
                    if isinstance(before_spacing, (int, float)):
                        paragraph.paragraph_format.space_before = Pt(before_spacing)
                        app_logger.debug(f"设置段前间距: {before_spacing}磅")
                    else:
                        app_logger.warning(f"段前间距值类型错误: {type(before_spacing)}")
            except Exception as e:
                app_logger.error(f"设置段前间距失败: {str(e)}")
            
            # 段后间距
            try:
                after_spacing = format_info.get('after_spacing', None)
                app_logger.debug(f"段后间距设置: {after_spacing}")
                
                if after_spacing is not None:
                    if isinstance(after_spacing, (int, float)):
                        paragraph.paragraph_format.space_after = Pt(after_spacing)
                        app_logger.debug(f"设置段后间距: {after_spacing}磅")
                    else:
                        app_logger.warning(f"段后间距值类型错误: {type(after_spacing)}")
            except Exception as e:
                app_logger.error(f"设置段后间距失败: {str(e)}")
                
        except Exception as e:
            app_logger.error(f"应用段落格式时发生异常: {str(e)}")
            # 不抛出异常，使用默认段落设置
    
    def get_output_file(self):
        """
        获取输出文件路径
        
        Returns:
            输出文件路径
        """
        return self.output_file
